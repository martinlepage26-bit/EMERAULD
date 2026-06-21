#!/usr/bin/env python3
"""Generate submission-ready .docx for Paper 25 — The Pharos Frame.
Target: AI & Society — Open Forum section.
Format: Springer submission standard — double-spaced, 12pt Times New Roman,
        numbered sections, APA/Harvard references, Appendix after references.
"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Source ─────────────────────────────────────────────────────────────────
SOURCE = "/mnt/c/Users/softinfo/Documents/EMERAULD/wiki/Paper 25 — The Pharos Frame (Draft 2026-04-23).md"
OUTPUT = "/mnt/c/Users/softinfo/Documents/EMERAULD/assets/Paper25_ThePharosFrame_AISociety_2026-04-26.docx"


# ── Document setup ──────────────────────────────────────────────────────────
def make_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    # Default Normal style
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    # Heading 1 (section)
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(12)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # Heading 2 (subsection)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.italic = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(3)
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    return doc


# ── Style helpers ───────────────────────────────────────────────────────────
def add_normal(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=None, first_line=False):
    """Add a paragraph with Normal style."""
    p = doc.add_paragraph(style='Normal')
    p.alignment = align
    if indent:
        p.paragraph_format.left_indent = indent
    if first_line:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_para_with_inline(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False):
    """Add a paragraph parsing inline **bold** and *italic* markdown."""
    p = doc.add_paragraph(style='Normal')
    p.alignment = align
    if first_line:
        p.paragraph_format.first_line_indent = Cm(1.25)
    _add_inline(p, text)
    return p


def _add_inline(p, text):
    """Parse inline markdown bold/italic and add runs to paragraph p."""
    # Pattern: **bold**, *italic*, ***bold-italic***
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)')
    pos = 0
    for m in pattern.finditer(text):
        # Text before match
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        full = m.group(0)
        if full.startswith('***'):
            run = p.add_run(m.group(2))
            run.bold = True; run.italic = True
        elif full.startswith('**'):
            run = p.add_run(m.group(3))
            run.bold = True
        else:
            run = p.add_run(m.group(4))
            run.italic = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def add_heading(doc, text, level):
    """Add section heading."""
    style = 'Heading 1' if level == 1 else 'Heading 2'
    p = doc.add_heading(text, level=level)
    p.style = doc.styles[style]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Clear default heading color artifacts
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_table_row(doc, cells, bold_first=False):
    """Not used directly — tables built inline."""
    pass


def set_cell_shading(cell, fill):
    """Helper: set table cell shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


# ── Title page ─────────────────────────────────────────────────────────────
def add_title_page(doc):
    # Journal target
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Submitted to: AI & Society — Open Forum")
    run.italic = True
    run.font.size = Pt(11)
    doc.add_paragraph(style='Normal')  # spacer

    # Title
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    run = p.add_run("The Pharos Frame: Four Levels Where Ethics Becomes AI Governance")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(style='Normal')  # spacer

    # Author
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Martin Lepage, PhD")

    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("PHAROS Inc., Montréal, Québec, Canada")

    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ml@pharos-ai.ca")
    run.italic = True

    # Page break after title page
    doc.add_page_break()


# ── Abstract block ──────────────────────────────────────────────────────────
def add_abstract(doc, abstract_text, keywords):
    add_heading(doc, "Abstract", 1)
    add_para_with_inline(doc, abstract_text)

    doc.add_paragraph(style='Normal')
    p = doc.add_paragraph(style='Normal')
    run = p.add_run("Keywords: ")
    run.bold = True
    p.add_run(keywords)
    doc.add_page_break()


# ── Table builder ──────────────────────────────────────────────────────────
def add_markdown_table(doc, table_md):
    """Parse a simple markdown table and insert a Word table."""
    lines = [l.strip() for l in table_md.strip().split('\n') if l.strip()]
    # Filter out separator lines (---|---|---)
    data_lines = [l for l in lines if not re.match(r'^[\|\s\-:]+$', l)]
    rows = []
    for line in data_lines:
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    ncols = len(rows[0])
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = tbl.cell(i, j)
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            _add_inline(p, cell_text)
            if i == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph(style='Normal')


# ── Main content parser ────────────────────────────────────────────────────
def parse_body(doc, md_text):
    """Parse markdown sections and write into doc."""
    lines = md_text.split('\n')
    i = 0
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]

        # Skip wiki-only markers
        if line.startswith('---') and i in (0, len(lines)-1):
            i += 1
            continue

        # Section headings
        if line.startswith('### '):
            if in_table:
                add_markdown_table(doc, '\n'.join(table_lines))
                table_lines = []; in_table = False
            add_heading(doc, line[4:].strip(), 2)
            i += 1
            continue

        if line.startswith('## '):
            if in_table:
                add_markdown_table(doc, '\n'.join(table_lines))
                table_lines = []; in_table = False
            heading_text = line[3:].strip()
            # Skip wiki-only sections
            if heading_text in ('Related',):
                break
            add_heading(doc, heading_text, 1)
            i += 1
            continue

        # Table rows
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table:
                add_markdown_table(doc, '\n'.join(table_lines))
                table_lines = []; in_table = False

        # Empty line — paragraph break
        if not line.strip():
            i += 1
            continue

        # Horizontal rule — page break before major sections
        if re.match(r'^-{3,}$', line):
            i += 1
            continue

        # Italic-only paragraph (note / subhead)
        if line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            p = doc.add_paragraph(style='Normal')
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(line.strip('*'))
            run.italic = True
            i += 1
            continue

        # Bold paragraph acting as subheading (**text**)
        if line.startswith('**') and line.endswith('**') and line.count('**') == 2:
            p = doc.add_paragraph(style='Normal')
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(line.strip('*'))
            run.bold = True
            i += 1
            continue

        # Normal paragraph — collect continuation lines
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and not lines[i].startswith('|') and not lines[i].startswith('---'):
            para_lines.append(lines[i])
            i += 1
        full_text = ' '.join(para_lines)
        # Strip trailing draft note (the italic block after title)
        if full_text.startswith('*Draft ') or full_text.startswith('*Martin Lepage'):
            continue
        add_para_with_inline(doc, full_text, first_line=True)
        continue

    if in_table:
        add_markdown_table(doc, '\n'.join(table_lines))


# ── Main ───────────────────────────────────────────────────────────────────
def build():
    with open(SOURCE, 'r', encoding='utf-8') as f:
        raw = f.read()

    # Strip YAML frontmatter
    raw = re.sub(r'^---.*?---\n', '', raw, flags=re.DOTALL)

    # Strip wiki-only trailing sections
    raw = re.sub(r'\n## Related.*', '', raw, flags=re.DOTALL)

    # Separate abstract
    abstract_match = re.search(r'## Abstract\n\n(.*?)\n\n\*\*Keywords:', raw, re.DOTALL)
    abstract_text = abstract_match.group(1).strip() if abstract_match else ""
    keywords_match = re.search(r'\*\*Keywords:\*\*\s*(.*?)\n', raw)
    keywords = keywords_match.group(1).strip() if keywords_match else ""

    # Separate body: from Introduction onward, through References and Appendix
    body_start = raw.find('\n## 1.')
    if body_start < 0:
        body_start = raw.find('\n## 1 ')
    body_text = raw[body_start:].strip() if body_start >= 0 else ""

    doc = make_doc()
    add_title_page(doc)
    add_abstract(doc, abstract_text, keywords)
    parse_body(doc, body_text)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == '__main__':
    build()
