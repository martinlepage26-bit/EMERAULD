#!/usr/bin/env python3
"""Final formatting pass: Times New Roman 12pt, single-spaced, no orphans."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class DocumentFormatter:
    def __init__(self, bundle_dir):
        self.bundle_dir = Path(bundle_dir)

    def read_file(self, filename):
        """Read a file and return its content."""
        filepath = self.bundle_dir / filename
        try:
            if filepath.suffix == '.docx':
                doc = Document(str(filepath))
                return '\n\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            else:
                with open(filepath, 'r', encoding='utf-8') as f:
                    return f.read()
        except Exception as e:
            print(f"[ERROR] Error reading {filename}: {e}")
            return ""

    def clean_provenance_metadata(self, content):
        """Remove draft provenance metadata sections."""
        content = re.sub(
            r'## Provenance Metadata\n.*?(?=\n---|\n##)',
            '',
            content,
            flags=re.DOTALL
        )
        content = re.sub(
            r'## Source Note\n.*?(?=\n---|\n##|\n#)',
            '',
            content,
            flags=re.DOTALL
        )
        return content

    def standardize_date_placeholders(self, content):
        """Replace date placeholders."""
        content = re.sub(
            r'\[DATE\s*\+?\s*\d*\s*(?:months?|days?|weeks?)?\]',
            '[DATE TO BE FILLED]',
            content,
            flags=re.IGNORECASE
        )
        return content

    def set_paragraph_formatting(self, para):
        """Apply standard formatting to paragraph."""
        # Font: Times New Roman, 12pt
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Line spacing: single (1.0)
        para_format = para.paragraph_format
        para_format.line_spacing = 1.0

        # Space after paragraph
        para_format.space_after = Pt(6)

        # Widow/orphan control
        pPr = para._element.get_or_add_pPr()
        widowControl = OxmlElement('w:widowControl')
        pPr.append(widowControl)

    def add_paragraph_with_formatting(self, doc, text, style='Normal', is_heading=False, heading_level=None):
        """Add paragraph with proper formatting."""
        if not text.strip():
            doc.add_paragraph()
            return

        if is_heading:
            para = doc.add_heading(text.strip(), level=heading_level)
        else:
            para = doc.add_paragraph(text.strip(), style=style)

        # Apply formatting
        self.set_paragraph_formatting(para)

        # For headings, keep them bold but ensure font settings
        if is_heading:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.bold = True

    def markdown_to_docx(self, markdown_content, output_filename):
        """Convert markdown content to .docx with final formatting."""
        doc = Document()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Set document author
        core_props = doc.core_properties
        core_props.author = 'Martin Lepage, PhD'

        lines = markdown_content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # Handle headings
            if line.startswith('# ') and not line.startswith('## '):
                self.add_paragraph_with_formatting(
                    doc, line[2:].strip(),
                    is_heading=True, heading_level=1
                )
                i += 1
            elif line.startswith('## ') and not line.startswith('### '):
                self.add_paragraph_with_formatting(
                    doc, line[3:].strip(),
                    is_heading=True, heading_level=2
                )
                i += 1
            elif line.startswith('### ') and not line.startswith('#### '):
                self.add_paragraph_with_formatting(
                    doc, line[4:].strip(),
                    is_heading=True, heading_level=3
                )
                i += 1
            elif line.startswith('#### '):
                self.add_paragraph_with_formatting(
                    doc, line[5:].strip(),
                    is_heading=True, heading_level=4
                )
                i += 1
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet list
                bullet_text = line[2:].strip()
                para = doc.add_paragraph(bullet_text, style='List Bullet')
                self.set_paragraph_formatting(para)
                i += 1
            elif line.startswith('| '):
                # Table
                table_lines = [line]
                i += 1
                while i < len(lines) and lines[i].startswith('| '):
                    table_lines.append(lines[i])
                    i += 1
                self._add_markdown_table(doc, table_lines)
            elif line.strip() == '':
                # Blank line - add empty paragraph for spacing
                para = doc.add_paragraph()
                self.set_paragraph_formatting(para)
                i += 1
            else:
                # Regular paragraph
                if line.strip():
                    self.add_paragraph_with_formatting(doc, line.strip(), style='Normal')
                i += 1

        doc.save(str(output_filename))
        return str(output_filename)

    def _add_markdown_table(self, doc, table_lines):
        """Parse and add a markdown table to the document."""
        if len(table_lines) < 2:
            return

        # Parse header
        header_line = table_lines[0]
        headers = [h.strip() for h in header_line.split('|')[1:-1]]

        # Create table
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'

        # Add header row
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            # Format header
            for paragraph in cell.paragraphs:
                self.set_paragraph_formatting(paragraph)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

        # Add data rows (skip separator line)
        for line in table_lines[2:]:
            if '|' not in line or line.strip().startswith('|-'):
                continue
            row_data = [cell.strip() for cell in line.split('|')[1:-1]]
            if row_data and any(row_data):
                row = table.add_row()
                for i, cell_data in enumerate(row_data):
                    if i < len(row.cells):
                        row.cells[i].text = cell_data
                        # Format table cells
                        for paragraph in row.cells[i].paragraphs:
                            self.set_paragraph_formatting(paragraph)
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(12)

    def process_document(self, title, source_files, base_file, output_name):
        """Process a document through revision and conversion."""
        print(f"Formatting: {title}...", end=' ', flush=True)

        # Read base version
        base_content = self.read_file(base_file)

        # Apply revisions
        revised_content = base_content

        if '## Provenance Metadata' in revised_content or '## Source Note' in revised_content:
            revised_content = self.clean_provenance_metadata(revised_content)

        if '[DATE' in revised_content:
            revised_content = self.standardize_date_placeholders(revised_content)

        # Remove excess blank lines
        revised_content = re.sub(r'\n\n\n+', '\n\n', revised_content)

        # Convert to .docx with final formatting
        output_path = self.bundle_dir / output_name
        self.markdown_to_docx(revised_content, output_path)

        print(f"✓ {output_path.stat().st_size / 1024:.0f} KB")
        return str(output_path)


def main():
    bundle_dir = Path(__file__).parent
    formatter = DocumentFormatter(bundle_dir)

    documents = [
        {
            'title': 'PHAROS Invention Disclosure V12',
            'sources': ['00_PHAROS_Invention_Disclosure_v12.docx'],
            'base': '00_PHAROS_Invention_Disclosure_v12.docx',
            'output': 'REV_PHAROS_Invention_Disclosure_v12.docx'
        },
        {
            'title': 'Recursive AI Governance Very Long Narrative',
            'sources': ['01_Recursive_AI_Governance_Very_Long_Narrative.md'],
            'base': '01_Recursive_AI_Governance_Very_Long_Narrative.md',
            'output': 'REV_Recursive_AI_Governance_Very_Long_Narrative.docx'
        },
        {
            'title': 'PHAROS Master SOP',
            'sources': ['02_PHAROS_Master_SOP.md'],
            'base': '02_PHAROS_Master_SOP.md',
            'output': 'REV_PHAROS_Master_SOP.docx'
        },
        {
            'title': 'PHAROS Redesign Passes (Claude+GPT)',
            'sources': ['03_PHAROS_Recursive_Redesign_Passes_1-5_Claude+GPT.txt'],
            'base': '03_PHAROS_Recursive_Redesign_Passes_1-5_Claude+GPT.txt',
            'output': 'REV_PHAROS_Recursive_Redesign_Passes_1-5_Claude+GPT.docx'
        },
        {
            'title': 'PHAROS Redesign Passes (Claude)',
            'sources': ['04_PHAROS_Recursive_Redesign_Passes_1-5_Claude.txt'],
            'base': '04_PHAROS_Recursive_Redesign_Passes_1-5_Claude.txt',
            'output': 'REV_PHAROS_Recursive_Redesign_Passes_1-5_Claude.docx'
        },
        {
            'title': 'From Recursive Production to Governable Method',
            'sources': ['05_From_Recursive_Production_to_Governable_Method.md'],
            'base': '05_From_Recursive_Production_to_Governable_Method.md',
            'output': 'REV_From_Recursive_Production_to_Governable_Method.docx'
        },
        {
            'title': 'Codex Governance Case Study',
            'sources': ['06_Codex_Governance_Case_Study_Blind_Leading_Automated.md'],
            'base': '06_Codex_Governance_Case_Study_Blind_Leading_Automated.md',
            'output': 'REV_Codex_Governance_Case_Study_Blind_Leading_Automated.docx'
        },
        {
            'title': 'InfraFabric Architecture',
            'sources': ['07_InfraFabric_Architecture.md'],
            'base': '07_InfraFabric_Architecture.md',
            'output': 'REV_InfraFabric_Architecture.docx'
        },
        {
            'title': 'The Möbius Protocol',
            'sources': ['08_Mobius_Protocol_Complete_Operational_Specification.md'],
            'base': '08_Mobius_Protocol_Complete_Operational_Specification.md',
            'output': 'REV_Mobius_Protocol.docx'
        }
    ]

    print("\nApplying final formatting to all documents...\n")
    for doc_spec in documents:
        formatter.process_document(
            doc_spec['title'],
            doc_spec['sources'],
            doc_spec['base'],
            doc_spec['output']
        )

    print("\n✓ All files formatted with Times New Roman 12pt, single-spaced, no orphans")
    print("✓ Author set to: Martin Lepage, PhD")


if __name__ == '__main__':
    main()
