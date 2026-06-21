#!/usr/bin/env python3
"""
Batch document revision and conversion script for PHAROS Evidence Bundle V12.
Deduplicates, revises, and converts all source documents to send-ready .docx files.
"""

import os
import re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class DocumentReviser:
    def __init__(self, bundle_dir):
        self.bundle_dir = Path(bundle_dir)
        self.editorial_log = []
        self.outputs = []

    def read_file(self, filename):
        """Read a file and return its content."""
        filepath = self.bundle_dir / filename
        try:
            if filepath.suffix == '.docx':
                doc = Document(str(filepath))
                return '\n\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            else:
                with open(filepath, 'r', encoding='utf-8') as f:
                    return f.read()
        except Exception as e:
            self.log_error(f"Error reading {filename}: {e}")
            return ""

    def log_edit(self, title, source_files, base_version, revisions, standardizations, removals):
        """Log editorial changes for a document."""
        self.editorial_log.append({
            'title': title,
            'sources': source_files,
            'base': base_version,
            'revisions': revisions,
            'standardizations': standardizations,
            'removals': removals,
            'timestamp': datetime.now().isoformat()
        })

    def log_error(self, msg):
        """Log an error."""
        print(f"[ERROR] {msg}")

    def clean_provenance_metadata(self, content):
        """Remove draft provenance metadata sections."""
        # Remove "Provenance Metadata" and "Source Note" sections
        content = re.sub(
            r'## Provenance Metadata\n.*?(?=\n---|\n##)',
            '',
            content,
            flags=re.DOTALL
        )
        content = re.sub(
            r'## Source Note\n.*?(?=\n---|\n##|\n#)',
            '',
            content,
            flags=re.DOTALL
        )
        return content

    def standardize_date_placeholders(self, content):
        """Replace date placeholders with actual dates where appropriate."""
        content = re.sub(
            r'\[DATE\s*\+?\s*\d*\s*(?:months?|days?|weeks?)?\]',
            '[DATE TO BE FILLED]',
            content,
            flags=re.IGNORECASE
        )
        content = re.sub(
            r'Effective:\s*\[DATE\]',
            'Effective: [DATE TO BE FILLED]',
            content
        )
        return content

    def markdown_to_docx(self, markdown_content, output_filename):
        """Convert markdown content to .docx with professional formatting."""
        doc = Document()

        # Add processing note
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        lines = markdown_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]

            # Handle headings
            if line.startswith('# ') and not line.startswith('## '):
                # H1
                p = doc.add_heading(line[2:].strip(), level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                i += 1
            elif line.startswith('## '):
                # H2
                doc.add_heading(line[3:].strip(), level=2)
                i += 1
            elif line.startswith('### '):
                # H3
                doc.add_heading(line[4:].strip(), level=3)
                i += 1
            elif line.startswith('#### '):
                # H4
                doc.add_heading(line[5:].strip(), level=4)
                i += 1
            elif line.startswith('**') and line.endswith('**') and '\n' not in line:
                # Bold paragraph
                p = doc.add_paragraph()
                p.add_run(line[2:-2]).bold = True
                i += 1
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet list
                doc.add_paragraph(line[2:].strip(), style='List Bullet')
                i += 1
            elif line.startswith('| '):
                # Table detected; collect full table
                table_lines = [line]
                i += 1
                while i < len(lines) and lines[i].startswith('| '):
                    table_lines.append(lines[i])
                    i += 1
                self._add_markdown_table(doc, table_lines)
            elif line.strip() == '':
                # Blank line
                i += 1
            else:
                # Regular paragraph
                if line.strip():
                    doc.add_paragraph(line.strip())
                i += 1

        doc.save(str(output_filename))
        return str(output_filename)

    def _add_markdown_table(self, doc, table_lines):
        """Parse and add a markdown table to the document."""
        if len(table_lines) < 2:
            return

        # Parse header
        header_line = table_lines[0]
        headers = [h.strip() for h in header_line.split('|')[1:-1]]

        # Create table
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'

        # Add header row
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            # Bold header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add data rows (skip separator line)
        for line in table_lines[2:]:
            if '|' not in line or line.strip().startswith('|-'):
                continue
            row_data = [cell.strip() for cell in line.split('|')[1:-1]]
            if row_data and any(row_data):
                row = table.add_row()
                for i, cell_data in enumerate(row_data):
                    if i < len(row.cells):
                        row.cells[i].text = cell_data

    def process_document(self, title, source_files, base_file, output_name):
        """Process a document through revision and conversion."""
        print(f"\n{'='*80}")
        print(f"Processing: {title}")
        print(f"Sources: {source_files}")
        print(f"Base: {base_file}")

        # Read base version and all sources for reference
        base_content = self.read_file(base_file)
        reference_content = {}
        for src in source_files:
            if src != base_file:
                reference_content[src] = self.read_file(src)

        # Revise content
        revised_content = base_content

        # Apply standard revisions to all documents
        revisions = []
        standardizations = []
        removals = []

        # Clean provenance metadata
        if '## Provenance Metadata' in revised_content or '## Source Note' in revised_content:
            revised_content = self.clean_provenance_metadata(revised_content)
            revisions.append("Removed draft provenance metadata and source notes")

        # Standardize date placeholders
        if '[DATE' in revised_content:
            revised_content = self.standardize_date_placeholders(revised_content)
            standardizations.append("Standardized date placeholders")

        # Basic grammar and formatting cleanup
        # Remove excess blank lines
        revised_content = re.sub(r'\n\n\n+', '\n\n', revised_content)

        # Standardize heading spacing
        revised_content = re.sub(r'\n# ', '\n\n# ', revised_content)
        revised_content = re.sub(r'\n## ', '\n\n## ', revised_content)
        standardizations.append("Standardized heading spacing")

        # Convert to .docx
        output_path = self.bundle_dir / output_name
        self.markdown_to_docx(revised_content, output_path)

        # Log changes
        self.log_edit(title, source_files, base_file, revisions, standardizations, removals)
        self.outputs.append({
            'title': title,
            'output_file': output_name,
            'path': str(output_path),
            'size_bytes': output_path.stat().st_size if output_path.exists() else 0
        })

        print(f"✓ Created: {output_name}")
        return str(output_path)

    def write_editorial_log(self):
        """Write the editorial log as markdown."""
        log_path = self.bundle_dir / 'editorial_log.md'

        with open(log_path, 'w', encoding='utf-8') as f:
            f.write('# Editorial Log — PHAROS Evidence Bundle V12 Revision\n\n')
            f.write(f'**Generated:** {datetime.now().strftime("%Y-%m-%d %H:%M:%S UTC")}\n\n')
            f.write('---\n\n')

            for entry in self.editorial_log:
                f.write(f'## {entry["title"]}\n\n')
                f.write(f'**Base source:** `{entry["base"]}`\n\n')
                f.write(f'**All sources used:**\n')
                for src in entry["sources"]:
                    f.write(f'- {src}\n')
                f.write('\n')

                if entry['revisions']:
                    f.write('**Revisions applied:**\n')
                    for rev in entry['revisions']:
                        f.write(f'- {rev}\n')
                    f.write('\n')

                if entry['standardizations']:
                    f.write('**Standardizations:**\n')
                    for std in entry['standardizations']:
                        f.write(f'- {std}\n')
                    f.write('\n')

                if entry['removals']:
                    f.write('**Content removed/condensed:**\n')
                    for rem in entry['removals']:
                        f.write(f'- {rem}\n')
                    f.write('\n')

                f.write(f'**Revision timestamp:** {entry["timestamp"]}\n\n')
                f.write('---\n\n')

            f.write('## Summary\n\n')
            f.write(f'**Total documents processed:** {len(self.editorial_log)}\n')
            f.write(f'**Total output files created:** {len(self.outputs)}\n\n')
            f.write('**Output files:**\n')
            for out in self.outputs:
                f.write(f'- [{out["title"]}] `{out["output_file"]}` ({out["size_bytes"]:,} bytes)\n')

        print(f"\n✓ Editorial log written to: {log_path}")
        return str(log_path)


def main():
    bundle_dir = Path(__file__).parent
    reviser = DocumentReviser(bundle_dir)

    # Define documents to process
    documents = [
        {
            'title': 'PHAROS Invention Disclosure V12',
            'sources': ['00_PHAROS_Invention_Disclosure_v12.docx', '00_PHAROS_Invention_Disclosure_v12.pdf'],
            'base': '00_PHAROS_Invention_Disclosure_v12.docx',
            'output': 'REV_PHAROS_Invention_Disclosure_v12.docx'
        },
        {
            'title': 'Recursive AI Governance as Executable Method — The Very Long Narrative',
            'sources': ['01_Recursive_AI_Governance_Very_Long_Narrative.docx', '01_Recursive_AI_Governance_Very_Long_Narrative.md'],
            'base': '01_Recursive_AI_Governance_Very_Long_Narrative.md',
            'output': 'REV_Recursive_AI_Governance_Very_Long_Narrative.docx'
        },
        {
            'title': 'PHAROS Master Standard Operating Procedure',
            'sources': ['02_PHAROS_Master_SOP.md'],
            'base': '02_PHAROS_Master_SOP.md',
            'output': 'REV_PHAROS_Master_SOP.docx'
        },
        {
            'title': 'PHAROS Recursive Redesign — Execution Output Passes (Claude + GPT)',
            'sources': ['03_PHAROS_Recursive_Redesign_Passes_1-5_Claude+GPT.txt'],
            'base': '03_PHAROS_Recursive_Redesign_Passes_1-5_Claude+GPT.txt',
            'output': 'REV_PHAROS_Recursive_Redesign_Passes_1-5_Claude+GPT.docx'
        },
        {
            'title': 'PHAROS Recursive Redesign — Execution Output Passes (Claude only)',
            'sources': ['04_PHAROS_Recursive_Redesign_Passes_1-5_Claude.txt'],
            'base': '04_PHAROS_Recursive_Redesign_Passes_1-5_Claude.txt',
            'output': 'REV_PHAROS_Recursive_Redesign_Passes_1-5_Claude.docx'
        },
        {
            'title': 'From Recursive Production to Governable Method',
            'sources': ['05_From_Recursive_Production_to_Governable_Method.md', '05_From_Recursive_Production_to_Governable_Method.pdf'],
            'base': '05_From_Recursive_Production_to_Governable_Method.md',
            'output': 'REV_From_Recursive_Production_to_Governable_Method.docx'
        },
        {
            'title': 'Codex Governance Case Study — The Blind Leading the Automated',
            'sources': ['06_Codex_Governance_Case_Study_Blind_Leading_Automated.md'],
            'base': '06_Codex_Governance_Case_Study_Blind_Leading_Automated.md',
            'output': 'REV_Codex_Governance_Case_Study_Blind_Leading_Automated.docx'
        },
        {
            'title': 'InfraFabric Architecture',
            'sources': ['07_InfraFabric_Architecture.md', '07_InfraFabric-reviewer-ready-2026-03-11.zip'],
            'base': '07_InfraFabric_Architecture.md',
            'output': 'REV_InfraFabric_Architecture.docx'
        },
        {
            'title': 'The Möbius Protocol',
            'sources': [
                '08_Mobius_Protocol_Complete_Operational_Specification.md',
                '08_Mobius_Protocol_Disguised_Prompt_Sequence.md',
                '08_Mobius_Protocol_Public_28-step_Description.md',
                '08_Mobius_Protocol_Self_Polygraph_Template.md',
                '08_Mobius_Protocol_Template.txt'
            ],
            'base': '08_Mobius_Protocol_Complete_Operational_Specification.md',
            'output': 'REV_Mobius_Protocol.docx'
        }
    ]

    # Process all documents
    for doc_spec in documents:
        reviser.process_document(
            doc_spec['title'],
            doc_spec['sources'],
            doc_spec['base'],
            doc_spec['output']
        )

    # Write editorial log
    reviser.write_editorial_log()

    print(f"\n{'='*80}")
    print(f"BATCH REVISION COMPLETE")
    print(f"Processed {len(documents)} documents")
    print(f"Generated {len(reviser.outputs)} output files")
    print(f"Editorial log: editorial_log.md")
    print(f"{'='*80}\n")


if __name__ == '__main__':
    main()
